#!/usr/bin/env python3
"""
Integration tests for document text extraction.
Tests real document processing across all supported formats.
"""
import sys
import tempfile
from pathlib import Path
from typing import List, Tuple

# Add parent to path
sys.path.insert(0, str(Path(__file__).parent))

from services.text_extraction import text_extraction
from services.text_validator import text_validator


class DocumentSampleTests:
    """Test extraction with real and synthetic documents."""
    
    def __init__(self):
        """Initialize test suite."""
        self.results = []
        self.temp_dir = Path(tempfile.gettempdir())
    
    # ========================================================================
    # TXT Format Tests
    # ========================================================================
    
    def create_sample_txt(self) -> str:
        """Create a sample TXT file."""
        content = """Real Estate Property Analysis Report

Property Address: 123 Oak Street, San Francisco, CA 94102
Assessment Date: 2024-01-15

EXECUTIVE SUMMARY
This property report provides a comprehensive analysis of the residential property
located at the above address. The assessment includes structural integrity evaluation,
material composition analysis, and hazard identification.

MATERIALS ASSESSMENT
- Foundation: Concrete, poured 1987, no visible cracks
- Roofing: Asphalt shingles, estimated 12 years old, fair condition
- Walls: Wood framing with brick veneer
- Windows: Original double-pane, thermal loss detected

CONSTRUCTION QUALITY ANALYSIS
The building demonstrates standard construction practices from the 1980s era.
Some cost-cutting measures were observed in the HVAC system installation,
potentially reducing efficiency by 15-20%.

RECOMMENDATIONS
1. Schedule professional roof inspection
2. Upgrade to energy-efficient windows
3. Improve attic insulation
4. Inspect for asbestos in pipe insulation

Assessment Value: $850,000
Confidence Level: High"""
        
        file_path = self.temp_dir / "test_sample.txt"
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return str(file_path)
    
    def test_txt_extraction(self) -> Tuple[bool, str]:
        """Test TXT file extraction."""
        try:
            print("\n" + "=" * 60)
            print("Testing TXT Format Extraction")
            print("=" * 60)
            
            # Create sample file
            file_path = self.create_sample_txt()
            print(f"✓ Created sample TXT file: {file_path}")
            
            # Extract text
            text, status, error = text_extraction.extract_text(file_path, 'txt')
            
            if error:
                print(f"✗ Extraction failed: {error}")
                return False, error
            
            print(f"✓ Extraction status: {status}")
            
            # Validate extracted text
            is_valid, sanitized, warning = text_validator.validate_and_sanitize(text)
            
            if not is_valid:
                print(f"✗ Text validation failed: {warning}")
                return False, warning
            
            print(f"✓ Text validation passed")
            
            # Get health score
            health = text_validator.get_text_health_score(sanitized)
            print(f"✓ Health score: {health['score']}/100")
            print(f"✓ Word count: {health['word_count']}")
            print(f"✓ Character count: {health['length']}")
            
            if warning:
                print(f"⚠️  Warning: {warning}")
            
            if health['warnings']:
                for w in health['warnings']:
                    print(f"⚠️  {w}")
            
            # Clean up
            Path(file_path).unlink()
            
            return True, "TXT extraction successful"
            
        except Exception as e:
            print(f"✗ Test failed: {str(e)}")
            return False, str(e)
    
    # ========================================================================
    # PDF Format Tests (if pdfplumber available)
    # ========================================================================
    
    def create_sample_pdf(self) -> str:
        """Create a sample PDF file."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
        except ImportError:
            print("⚠️  reportlab not installed, skipping PDF creation")
            return None
        
        file_path = self.temp_dir / "test_sample.pdf"
        
        c = canvas.Canvas(str(file_path), pagesize=letter)
        width, height = letter
        
        # Add content
        y = height - 50
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, y, "Property Inspection Report")
        
        y -= 30
        c.setFont("Helvetica", 12)
        
        text_lines = [
            "Property: 456 Elm Street, Oakland, CA 94601",
            "Inspector: John Smith, Licensed Inspector #12345",
            "",
            "STRUCTURAL FINDINGS:",
            "- Foundation: Concrete slab, 1992",
            "- Framing: Conventional wood frame",
            "- Exterior: Fiber cement siding",
            "",
            "HAZARDOUS MATERIALS DETECTED:",
            "- Lead paint in older bedrooms (pre-1978)",
            "- Asbestos in pipe insulation",
            "",
            "COST-CUTTING ASSESSMENT:",
            "Evidence of cost reduction in HVAC routing,",
            "potentially reducing system efficiency.",
        ]
        
        for line in text_lines:
            c.drawString(50, y, line)
            y -= 15
        
        c.save()
        return str(file_path)
    
    def test_pdf_extraction(self) -> Tuple[bool, str]:
        """Test PDF file extraction."""
        try:
            print("\n" + "=" * 60)
            print("Testing PDF Format Extraction")
            print("=" * 60)
            
            # Create sample file
            file_path = self.create_sample_pdf()
            if not file_path:
                return True, "PDF test skipped - reportlab not installed"
            
            print(f"✓ Created sample PDF file: {file_path}")
            
            # Extract text
            text, status, error = text_extraction.extract_text(file_path, 'pdf')
            
            if error:
                print(f"✗ Extraction failed: {error}")
                return False, error
            
            print(f"✓ Extraction status: {status}")
            
            # Validate extracted text
            is_valid, sanitized, warning = text_validator.validate_and_sanitize(text)
            
            if not is_valid:
                print(f"✗ Text validation failed: {warning}")
                return False, warning
            
            print(f"✓ Text validation passed")
            
            # Get health score
            health = text_validator.get_text_health_score(sanitized)
            print(f"✓ Health score: {health['score']}/100")
            print(f"✓ Word count: {health['word_count']}")
            
            # Clean up
            if Path(file_path).exists():
                Path(file_path).unlink()
            
            return True, "PDF extraction successful"
            
        except Exception as e:
            print(f"✗ Test failed: {str(e)}")
            return False, str(e)
    
    # ========================================================================
    # DOCX Format Tests
    # ========================================================================
    
    def create_sample_docx(self) -> str:
        """Create a sample DOCX file."""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError:
            print("⚠️  python-docx not installed, skipping DOCX creation")
            return None
        
        file_path = self.temp_dir / "test_sample.docx"
        doc = Document()
        
        # Add heading
        heading = doc.add_heading('Property Valuation Report', 0)
        
        # Add paragraphs
        doc.add_paragraph('Property Address: 789 Pine Street, Los Angeles, CA 90001')
        doc.add_paragraph('Valuation Date: January 15, 2024')
        
        doc.add_heading('Market Analysis', level=1)
        doc.add_paragraph(
            'The subject property is located in a desirable neighborhood with strong '
            'appreciation trends. Comparable properties have sold for 12-15% above '
            'list price in the past 6 months.'
        )
        
        doc.add_heading('Material Assessment', level=1)
        doc.add_paragraph(
            'The property features quality construction materials. Foundation is concrete, '
            'walls are brick veneer over wood frame, and roofing is architectural shingles. '
            'Some cost-cutting observed in plumbing sizing.'
        )
        
        # Add table
        doc.add_heading('Property Specifications', level=2)
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.cell(0, 0).text = 'Feature'
        table.cell(0, 1).text = 'Specification'
        table.cell(1, 0).text = 'Square Footage'
        table.cell(1, 1).text = '2,450 sq ft'
        table.cell(2, 0).text = 'Lot Size'
        table.cell(2, 1).text = '0.35 acres'
        table.cell(3, 0).text = 'Year Built'
        table.cell(3, 1).text = '1998'
        
        doc.save(file_path)
        return str(file_path)
    
    def test_docx_extraction(self) -> Tuple[bool, str]:
        """Test DOCX file extraction."""
        try:
            print("\n" + "=" * 60)
            print("Testing DOCX Format Extraction")
            print("=" * 60)
            
            # Create sample file
            file_path = self.create_sample_docx()
            if not file_path:
                return True, "DOCX test skipped - python-docx not installed"
            
            print(f"✓ Created sample DOCX file: {file_path}")
            
            # Extract text
            text, status, error = text_extraction.extract_text(file_path, 'docx')
            
            if error:
                print(f"✗ Extraction failed: {error}")
                return False, error
            
            print(f"✓ Extraction status: {status}")
            
            # Validate extracted text
            is_valid, sanitized, warning = text_validator.validate_and_sanitize(text)
            
            if not is_valid:
                print(f"✗ Text validation failed: {warning}")
                return False, warning
            
            print(f"✓ Text validation passed")
            
            # Get health score
            health = text_validator.get_text_health_score(sanitized)
            print(f"✓ Health score: {health['score']}/100")
            print(f"✓ Word count: {health['word_count']}")
            
            # Verify table content extracted
            if 'Square Footage' in sanitized or '2,450' in sanitized:
                print(f"✓ Table content extracted successfully")
            
            # Clean up
            if Path(file_path).exists():
                Path(file_path).unlink()
            
            return True, "DOCX extraction successful"
            
        except Exception as e:
            print(f"✗ Test failed: {str(e)}")
            return False, str(e)
    
    def run_all_tests(self):
        """Run all integration tests."""
        print("\n" + "=" * 60)
        print("DOCUMENT TEXT EXTRACTION INTEGRATION TESTS")
        print("=" * 60)
        
        results = []
        
        # TXT tests
        passed, message = self.test_txt_extraction()
        results.append(("TXT Extraction", passed, message))
        
        # PDF tests
        passed, message = self.test_pdf_extraction()
        results.append(("PDF Extraction", passed, message))
        
        # DOCX tests
        passed, message = self.test_docx_extraction()
        results.append(("DOCX Extraction", passed, message))
        
        # Print summary
        print("\n" + "=" * 60)
        print("TEST SUMMARY")
        print("=" * 60)
        
        passed_count = 0
        for test_name, passed, message in results:
            status = "✓ PASSED" if passed else "✗ FAILED"
            print(f"{status}: {test_name}")
            if not passed or "skipped" in message.lower():
                print(f"       {message}")
            passed_count += 1 if passed else 0
        
        print(f"\nTotal: {passed_count}/{len(results)} tests passed")
        print("=" * 60)
        
        return all(result[1] for result in results)


if __name__ == "__main__":
    test_suite = DocumentSampleTests()
    success = test_suite.run_all_tests()
    sys.exit(0 if success else 1)
