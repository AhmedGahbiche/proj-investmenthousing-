"""
Text extraction service module for extracting content from various document formats.
Supports PDF, DOCX, PNG (with OCR), and TXT formats.
Includes timeout protection, error handling, and text validation.
"""
import logging
import signal
from pathlib import Path
from typing import Tuple
from functools import wraps
from services.text_validator import text_validator

logger = logging.getLogger(__name__)


def timeout_handler(signum, frame):
    """Handle timeout during text extraction."""
    raise TimeoutError("Text extraction operation timed out")


def extract_with_timeout(timeout_seconds: int = 60):
    """
    Decorator to add timeout protection to extraction methods.
    
    Args:
        timeout_seconds: Maximum time allowed for extraction
    """
    def decorator(func):
        @wraps(func)
        def wrapper(*args, **kwargs):
            # Set signal handler
            signal.signal(signal.SIGALRM, timeout_handler)
            signal.alarm(timeout_seconds)
            try:
                result = func(*args, **kwargs)
                signal.alarm(0)  # Cancel alarm
                return result
            except TimeoutError:
                logger.error(f"Extraction timeout after {timeout_seconds}s")
                raise
            finally:
                signal.alarm(0)  # Ensure alarm is cancelled
        return wrapper
    return decorator


class TextExtractionService:
    """Service for extracting text from various document formats."""
    
    SUPPORTED_FORMATS = {
        'pdf': 'extract_text_from_pdf',
        'docx': 'extract_text_from_docx',
        'png': 'extract_text_from_png',
        'txt': 'extract_text_from_txt',
    }
    
    # Extraction timeout configuration (seconds)
    EXTRACTION_TIMEOUTS = {
        'pdf': 120,    # PDFs can be large
        'docx': 60,    # DOCX is usually smaller
        'png': 90,     # OCR can be slow
        'txt': 30,     # Text files should be fast
    }
    
    def __init__(self):
        """Initialize text extraction service."""
        self.validator = text_validator
    
    def extract_text(self, file_path: str, file_format: str) -> Tuple[str, str, str]:
        """
        Extract text from document based on format.
        
        Args:
            file_path: Path to the file
            file_format: File format (pdf, docx, png, txt)
            
        Returns:
            Tuple of (extracted_text, status, error_message)
            status can be: 'success', 'partial', 'failed'
            error_message is None if extraction succeeds
        """
        file_format = file_format.lower()
        
        if file_format not in self.SUPPORTED_FORMATS:
            error_msg = f"Unsupported file format: {file_format}"
            logger.error(error_msg)
            return "", "failed", error_msg
        
        try:
            # Get extraction method
            extraction_method = getattr(self, self.SUPPORTED_FORMATS[file_format])
            
            # Get timeout for this format
            timeout = self.EXTRACTION_TIMEOUTS.get(file_format, 60)
            
            # Extract text
            text, status = extraction_method(file_path, timeout)
            
            # Validate and sanitize extracted text
            is_valid, sanitized_text, validation_warning = self.validator.validate_and_sanitize(text)
            
            if validation_warning:
                logger.warning(f"Text validation warning: {validation_warning}")
            
            if not is_valid:
                logger.warning(f"Text validation failed: {validation_warning}")
                status = "failed" if status != "success" else "partial"
                return "", status, validation_warning
            
            logger.info(f"Text extraction completed for {file_format}: {status}")
            return sanitized_text, status, None
            
        except TimeoutError as e:
            error_msg = f"Extraction timeout for {file_format}: {str(e)}"
            logger.error(error_msg)
            return "", "failed", error_msg
        except Exception as e:
            error_msg = f"Error extracting text from {file_format}: {str(e)}"
            logger.error(error_msg)
            return "", "failed", error_msg
    
    @staticmethod
    @extract_with_timeout(timeout_seconds=30)
    def extract_text_from_txt(file_path: str, timeout: int = 30) -> Tuple[str, str]:
        """
        Extract text from TXT file.
        
        Args:
            file_path: Path to TXT file
            timeout: Timeout in seconds (unused, decorator handles it)
            
        Returns:
            Tuple of (text, status)
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            if not text.strip():
                logger.warning(f"TXT file is empty: {file_path}")
                return "", "partial"
            
            logger.info(f"Successfully extracted text from TXT: {file_path}")
            return text, "success"
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as f:
                    text = f.read()
                logger.info(f"Successfully extracted text from TXT with fallback encoding: {file_path}")
                return text, "success"
            except Exception as e:
                logger.error(f"Failed to extract text from TXT: {str(e)}")
                raise
        except Exception as e:
            logger.error(f"Failed to read TXT file: {str(e)}")
            raise
    
    @staticmethod
    @extract_with_timeout(timeout_seconds=120)
    def extract_text_from_pdf(file_path: str, timeout: int = 120) -> Tuple[str, str]:
        """
        Extract text from PDF file using pdfplumber (better than PyPDF2).
        
        Args:
            file_path: Path to PDF file
            timeout: Timeout in seconds (unused, decorator handles it)
            
        Returns:
            Tuple of (text, status)
        """
        try:
            import pdfplumber
        except ImportError:
            logger.warning("pdfplumber not installed, falling back to PyPDF2")
            return TextExtractionService._extract_text_from_pdf_pypdf2(file_path, timeout)
        
        try:
            text_content = []
            
            with pdfplumber.open(file_path) as pdf:
                if len(pdf.pages) == 0:
                    logger.warning(f"PDF has no pages: {file_path}")
                    return "", "partial"
                
                for page_num, page in enumerate(pdf.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(f"--- Page {page_num} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {str(e)}")
            
            merged_text = "\n\n".join(text_content)
            
            if not merged_text.strip():
                logger.warning(f"No text could be extracted from PDF: {file_path}")
                return "", "partial"
            
            logger.info(f"Successfully extracted text from PDF: {file_path}")
            return merged_text, "success"
            
        except Exception as e:
            logger.error(f"Failed to extract text from PDF with pdfplumber: {str(e)}")
            raise
    
    @staticmethod
    def _extract_text_from_pdf_pypdf2(file_path: str, timeout: int = 120) -> Tuple[str, str]:
        """
        Fallback PDF extraction using PyPDF2.
        
        Args:
            file_path: Path to PDF file
            timeout: Timeout in seconds (informational)
            
        Returns:
            Tuple of (text, status)
        """
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("PyPDF2 is not installed")
        
        try:
            text_content = []
            
            with open(file_path, 'rb') as f:
                pdf_reader = PdfReader(f)
                
                if len(pdf_reader.pages) == 0:
                    logger.warning(f"PDF has no pages: {file_path}")
                    return "", "partial"
                
                for page_num, page in enumerate(pdf_reader.pages, 1):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append(f"--- Page {page_num} ---\n{page_text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract text from page {page_num}: {str(e)}")
            
            merged_text = "\n\n".join(text_content)
            
            if not merged_text.strip():
                logger.warning(f"No text could be extracted from PDF: {file_path}")
                return "", "partial"
            
            logger.info(f"Successfully extracted text from PDF with PyPDF2: {file_path}")
            return merged_text, "success"
            
        except Exception as e:
            logger.error(f"Failed to extract text from PDF with PyPDF2: {str(e)}")
            raise
    
    @staticmethod
    @extract_with_timeout(timeout_seconds=60)
    def extract_text_from_docx(file_path: str, timeout: int = 60) -> Tuple[str, str]:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
            timeout: Timeout in seconds (unused, decorator handles it)
            
        Returns:
            Tuple of (text, status)
        """
        try:
            from docx import Document
        except ImportError:
            raise ImportError("python-docx is not installed")
        
        try:
            doc = Document(file_path)
            
            text_content = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_content.append(" | ".join(row_text))
            
            merged_text = "\n".join(text_content)
            
            if not merged_text.strip():
                logger.warning(f"No text could be extracted from DOCX: {file_path}")
                return "", "partial"
            
            logger.info(f"Successfully extracted text from DOCX: {file_path}")
            return merged_text, "success"
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {str(e)}")
            raise
    
    @staticmethod
    def _preprocess_image(image):
        """
        Preprocess image to improve OCR accuracy.
        
        Args:
            image: PIL Image object
            
        Returns:
            Preprocessed PIL Image object
        """
        try:
            from PIL import ImageEnhance, ImageFilter
        except ImportError:
            logger.warning("PIL enhancement not available, skipping preprocessing")
            return image
        
        try:
            # Convert to grayscale if color
            if image.mode != 'L':
                image = image.convert('L')
            
            # Enhance contrast
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(2)
            
            # Enhance sharpness
            enhancer = ImageEnhance.Sharpness(image)
            image = enhancer.enhance(2)
            
            # Apply slight blur to reduce noise
            image = image.filter(ImageFilter.MedianFilter(size=3))
            
            logger.debug(f"Image preprocessing completed: {image.size}")
            return image
        except Exception as e:
            logger.warning(f"Image preprocessing failed: {str(e)}, continuing without preprocessing")
            return image
    
    @staticmethod
    @extract_with_timeout(timeout_seconds=90)
    def extract_text_from_png(file_path: str, timeout: int = 90) -> Tuple[str, str]:
        """
        Extract text from PNG file using OCR (Tesseract).
        Includes image preprocessing to improve OCR accuracy.
        
        Args:
            file_path: Path to PNG file
            timeout: Timeout in seconds (unused, decorator handles it)
            
        Returns:
            Tuple of (text, status)
        """
        try:
            import pytesseract
            from PIL import Image
        except ImportError:
            raise ImportError("pytesseract and/or Pillow are not installed")
        
        try:
            # Check if Tesseract is installed
            try:
                import subprocess
                result = subprocess.run(['tesseract', '--version'], 
                                       capture_output=True, timeout=5)
                if result.returncode != 0:
                    raise RuntimeError("Tesseract not properly installed")
            except FileNotFoundError:
                raise RuntimeError("Tesseract-OCR is not installed on this system. "
                                  "Please install it: https://github.com/UB-Mannheim/tesseract/wiki")
            
            # Open image
            image = Image.open(file_path)
            logger.info(f"Loaded image: {image.size} {image.mode}")
            
            # Preprocess image for better OCR
            image = TextExtractionService._preprocess_image(image)
            
            # Extract text using OCR
            text = pytesseract.image_to_string(image)
            
            if not text or not text.strip():
                logger.warning(f"No text could be extracted from PNG via OCR: {file_path}")
                return "", "partial"
            
            logger.info(f"Successfully extracted text from PNG via OCR: {file_path}")
            return text, "success"
            
        except RuntimeError as e:
            logger.error(f"OCR system error: {str(e)}")
            raise
        except Exception as e:
            logger.error(f"Failed to extract text from PNG: {str(e)}")
            raise


# Global text extraction service instance
text_extraction = TextExtractionService()
